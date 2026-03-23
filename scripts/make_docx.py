from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# -----------------------------------------------
# CHANGE THIS to match your book name
book_name = "Чуварот"
# -----------------------------------------------

corrected_file = f"text/{book_name}_corrected.txt"
output_docx    = f"docx/{book_name}.docx"

doc = Document()

# --- Page setup (A4) ---
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Default font ---
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(14)

# --- Add Table of Contents field ---
toc_heading = doc.add_heading("Содржина", level=1)
toc_para = doc.add_paragraph()
run = toc_para.add_run()

fldChar_begin = OxmlElement("w:fldChar")
fldChar_begin.set(qn("w:fldCharType"), "begin")

instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

fldChar_end = OxmlElement("w:fldChar")
fldChar_end.set(qn("w:fldCharType"), "end")

run._r.append(fldChar_begin)
run._r.append(instrText)
run._r.append(fldChar_end)

doc.add_page_break()

# --- Read and add corrected text ---
with open(corrected_file, encoding="utf-8") as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()

    if not line:
        doc.add_paragraph("")
        continue

    # Heading 1 — all-caps lines or lines that look like chapter titles
    # Adjust this pattern to match your specific book's heading style
    if line.isupper() and len(line) > 3:
        doc.add_heading(line, level=1)

    # Heading 2 — lines ending with colon, or short bold-looking lines
    elif line.endswith(":") and len(line) < 80:
        doc.add_heading(line, level=2)

    # Normal paragraph
    else:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save(output_docx)
print(f"Saved: {output_docx}")
print("")
print("Next steps:")
print("  1. Open the .docx file in Microsoft Word")
print("  2. Press Ctrl+A to select everything")
print("  3. Press F9 to update the Table of Contents")
print("  4. Save the file again")